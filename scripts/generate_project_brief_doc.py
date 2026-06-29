"""Generate TRVRSE-PROJECT-BRIEF.docx from markdown source."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "TRVRSE-PROJECT-BRIEF.docx"

NAVY = RGBColor(10, 22, 40)
BLUE = RGBColor(45, 125, 210)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY if level == 1 else BLUE


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)


def build() -> None:
    doc = Document()

    # Title
    title = doc.add_heading("Trvrse by Nexxogen", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = NAVY

    sub = doc.add_paragraph("Project Brief & Reference Document")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = BLUE

    doc.add_paragraph(
        "Source: Claude planning conversations | Company: Nexxogen Limited | Product: Trvrse"
    )

    # Agent guide
    add_heading(doc, "AI / Developer Quick Guide", 1)
    add_para(
        doc,
        "Trvrse is a cross-border fintech wallet (not a bank). Users fund from Nigerian "
        "banks, auto-detect location for currency, convert at live rates, send to any global "
        "bank via Wise rails, and pay with virtual cards.",
    )
    add_heading(doc, "Already built (this repo)", 2)
    for item in [
        "Next.js 14 investor landing page with animations",
        "Live forex via ExchangeRate-API",
        "Waitlist API + Supabase schema",
        "Pitch deck PDF at public/trvrse-pitch-deck.pdf",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Priority order", 2)
    steps = [
        "CAC register Nexxogen Limited (LLC)",
        "Open Access Bank business account",
        "CBN PSSP application via legal counsel (Mary)",
        "Wise + Adyen sandbox APIs",
        "Build React Native mobile MVP (26 screens)",
        "Apply BOI YES + Tony Elumelu Foundation",
        "Seed raise after traction; deposit ₦200M CBN escrow",
        "Deploy landing to Vercel for investors",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"{i}. {s}")

    # Product
    add_heading(doc, "1. Product Vision", 1)
    add_para(doc, "Problem: Nigerians abroad face declined cards, hidden FX fees, fragmented apps.")
    add_para(
        doc,
        "Solution: One app — fund locally, auto-detect country, convert at live rates, "
        "send to any bank or pay with virtual card.",
    )
    add_para(doc, 'Tagline: "Pay Anywhere. Always."', bold=True)

    # Business model
    add_heading(doc, "2. Business Model", 1)
    add_bullet(doc, "0.5% margin on currency conversion and transfers")
    add_bullet(doc, "Flat fee for small domestic NG transfers under ~$10")
    add_bullet(doc, "Virtual card interchange ~0.3% (future)")
    add_bullet(doc, "Premium subscription ₦2,500/month (future)")
    add_para(
        doc,
        "Unit economics: 0.5% on $20M volume = $100K/year. Sustainable at 2–3M+ transactions "
        "or higher average ticket sizes.",
    )

    # Legal
    add_heading(doc, "3. Legal & Regulatory (Nigeria)", 1)
    add_bullet(doc, "Entity: Nexxogen Limited — Private LLC")
    add_bullet(doc, "License: PSSP + Super Agent + Bank partnership")
    add_bullet(doc, "Capital: ₦200,000,000 in CBN escrow (proof of capital)")
    add_bullet(doc, "Legal fees: ~₦5–10M for application and compliance docs")
    add_bullet(doc, "Data: Nigeria Data Protection Act 2023; register with NDPC")
    add_para(
        doc,
        "Cannot process real money at scale without CBN approval. Can build MVP and beta "
        "under Access Bank partnership while application is pending.",
    )

    # Partnerships
    add_heading(doc, "4. Partnerships", 1)
    add_bullet(doc, "Nigeria: Access Bank (primary) or GTBank")
    add_bullet(doc, "Global: Wise API for bank lookup, settlement, 80+ countries")
    add_bullet(doc, "Do NOT integrate every bank directly — use payment orchestration")
    add_bullet(doc, "Year 1 corridor: Nigeria ↔ UK, US, Canada")

    # Funding
    add_heading(doc, "5. Funding Strategy", 1)
    add_bullet(doc, "BOI YES Programme — up to ₦10M (legal fees)")
    add_bullet(doc, "Tony Elumelu Foundation — $5,000 grant (non-dilutive)")
    add_bullet(doc, "Antler Nigeria — $100K for 10% (needs co-founder)")
    add_bullet(doc, "Y Combinator — after MVP traction")
    add_bullet(doc, "Seed target: $200K–$500K after MVP")

    # Tech
    add_heading(doc, "6. Technical Stack", 1)
    add_bullet(doc, "Mobile (future): React Native + Expo, TypeScript")
    add_bullet(doc, "Landing: Next.js 14, Tailwind, Framer Motion, GSAP, Three.js")
    add_bullet(doc, "Hosting: Vercel (free) | DB: Supabase (free)")
    add_bullet(doc, "Forex: ExchangeRate-API → Wise API when approved")

    # Tiers
    add_heading(doc, "7. User Tier System", 1)
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["Tier", "KYC", "Daily Limit", "Features"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows = [
        ("Basic", "Email/phone", "$50", "Wallet-to-wallet"),
        ("Verified", "BVN/ID", "$500", "Domestic bank, conversion"),
        ("Premium", "Full KYC", "$5,000", "International, 1.5% fee"),
        ("Business", "CAC docs", "$50,000", "Bulk, API access"),
    ]
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    # Competitors
    add_heading(doc, "8. Competitors", 1)
    add_bullet(doc, "Wise — FX leader; no location auto-switch")
    add_bullet(doc, "Revolut — cards; not Nigeria-first")
    add_bullet(doc, "Chipper Cash — Africa remittance; no virtual card")
    add_bullet(doc, "Trvrse edge: location FX + virtual cards + direct transfers in one app")

    # Use cases
    add_heading(doc, "9. Use Cases", 1)
    add_bullet(doc, "Chioma — lands Toronto, converts NGN to CAD, pays with virtual card")
    add_bullet(doc, "Tunde — paid USD remotely, sends to family in Lagos")
    add_bullet(doc, "Zainab — global business, pays suppliers in GBP/USD")
    add_bullet(doc, "Adekunle — tuition from Nigeria to UK account")

    # Contacts
    add_heading(doc, "10. Contacts & Links", 1)
    add_bullet(doc, "Invest: invest@nexxogen.com")
    add_bullet(doc, "CAC: https://cac.gov.ng")
    add_bullet(doc, "BOI: https://www.boi.ng")
    add_bullet(doc, "TEF: https://www.tonyelumelufoundation.org/teep")
    add_bullet(doc, "Antler: https://www.antler.co/location/nigeria")
    add_bullet(doc, "Legal: Oluwatodimu Ige & Associates (via Mary)")

    doc.add_paragraph()
    add_para(
        doc,
        "Full markdown reference: docs/TRVRSE-PROJECT-BRIEF.md | "
        "AI guide: docs/TRVRSE-AGENT-GUIDE.md",
        bold=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
